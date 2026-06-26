#!/usr/bin/env python3
"""
generate_irb.py — 영남대학교병원(YUMC) IRB 제출 문서 3종 생성기 (standalone)

입력:  study_info.json  (references/study_info.schema.json 스키마 준수)
출력:  1) research_protocol.docx        — 연구계획서 본문 (13개 섹션)
       2) protocol_summary.docx         — 임상연구 계획서 요약 [서식 1]
       3) consent_waiver_checklist.docx — 연구대상자 동의 면제 점검표 [서식 41]

정책:
  - Single source of truth: 모든 문서는 study_info.json 한 파일에서 생성.
  - Citation Grounding: references 항목은 pmid 또는 doi 가 있는 것만 출력.
  - 본 스크립트는 .docx 생성까지만 담당. BRIA 등 IRB 시스템 자동 제출 없음.

사용법:
    python generate_irb.py study_info.json --outdir ./out
"""
import argparse, json, os, sys, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BURGUNDY = RGBColor(0x80, 0x1A, 0x2B)
FONT = "맑은 고딕"

def set_font(run, size=10, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = color
    # 한글 폰트 적용
    rpr = run._element.get_or_add_rPr()
    from docx.oxml.ns import qn
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT)

def h(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sizes = {1: 14, 2: 12, 3: 11}
    set_font(r, size=sizes.get(level, 11), bold=True,
             color=BURGUNDY if level == 1 else None)
    p.space_before = Pt(8)
    return p

def para(doc, text, size=10, bold=False, bullet=False):
    style = 'List Bullet' if bullet else None
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p

def load(path):
    with open(path, encoding='utf-8') as f:
        return json.load(f)

def grounded_refs(study):
    """PMID/DOI 동반 인용만 통과 (Citation Grounding)."""
    out, dropped = [], []
    for ref in study.get('references', []):
        if ref.get('pmid') or ref.get('doi') or ref.get('registry'):
            out.append(ref)
        else:
            dropped.append(ref.get('citation', '(no citation)'))
    return out, dropped

# ---------------------------------------------------------------- 본문
def build_protocol(study, outdir):
    doc = Document()
    pi = study['investigators']['pi']
    inst = study.get('institution', '영남대학교병원')
    dept = study.get('department', '순환기내과')

    # 표지
    t = doc.add_paragraph(); r = t.add_run(study['title']); set_font(r, 16, True, BURGUNDY)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for line in [
        f"연구 책임자: {inst} {dept} {pi['name']}",
        f"{pi.get('name_en','')} {pi.get('title','')}".strip(),
        "Division of Cardiology, Department of Internal Medicine",
        "Yeungnam University Hospital",
        f"E-mail: {pi.get('email','')}",
    ]:
        if line.strip():
            p = doc.add_paragraph(); rr = p.add_run(line); set_font(rr, 10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 목차
    h(doc, "Table of Contents", 2)
    toc = ["1. 연구계획서 요약","2. 연구 배경","3. 연구 목적","4. 연구 관리 구조",
           "5. 연구 참여에 따른 위험 및 이익","6. 윤리적 측면","7. 연구 설계 및 계획",
           "8. 수행","9. 데이터 관리","10. 자료의 분석","11. 보고서","12. 연구문서의 보관","13. 참고문헌"]
    for item in toc: para(doc, item, 10)
    doc.add_page_break()

    # 1. 요약
    h(doc, "1. 연구계획서 요약", 1)
    para(doc, study['objective'])
    para(doc, f"연구 디자인: {study['design']}")

    # 2. 배경
    h(doc, "2. 연구 배경", 1)
    if study.get('background'): para(doc, study['background'])
    else: para(doc, "(연구 배경 서술 필요)")

    # 3. 목적
    h(doc, "3. 연구 목적", 1)
    para(doc, study['objective'])

    # 4. 관리 구조
    h(doc, "4. 연구 관리 구조", 1)
    h(doc, f"4.1 연구 기관: {inst} {dept}", 3)
    h(doc, f"4.2 임상연구 기간: {study.get('period','IRB 승인일로부터 2년')}", 3)
    h(doc, "4.3 연구자 명단", 3)
    para(doc, f"연구 책임자: {pi.get('affiliation', inst+' '+dept)}, {pi['name']} {pi.get('title','')}".strip())
    for co in study['investigators'].get('co_investigators', []):
        para(doc, f"공동 연구자: {co.get('affiliation','')}, {co['name']} {co.get('role','')}".strip())
    h(doc, "4.4 연구 대상자 선정기준", 3)
    para(doc, "1) 선정 기준", bold=True)
    for c in study['inclusion_criteria']: para(doc, c, bullet=True)
    para(doc, "2) 제외 기준", bold=True)
    for c in study['exclusion_criteria']: para(doc, c, bullet=True)
    h(doc, "4.5 목표 피험자의 수 및 산출 근거", 3)
    para(doc, study.get('target_n', ''))
    h(doc, "4.6 연구책임자", 3)
    para(doc, "연구 책임자는 데이터 입력에 대한 관리 및 책임이 있으며, 본 임상연구가 적절히 수행·완료되도록 연구 내용을 숙지하고 의학적 판단 기준과 정보제공의 의무를 가진다. 연구의 시작과 종료에 관한 계획을 주도·결정하고 관련 규정을 준수한다.")

    # 5. 위험/이익
    h(doc, "5. 연구 참여에 따른 위험 및 이익", 1)
    para(doc, study.get('risk_benefit',
        "이미 시행한 검사 결과를 바탕으로 통계분석을 시행하는 후향적 의무기록 연구이므로 연구 참여에 따른 직접적인 위험 및 이익은 없다."))

    # 6. 윤리
    h(doc, "6. 윤리적 측면", 1)
    para(doc, f"본 연구는 만 {study.get('age_min',19)}세 이상의 성인을 대상으로 한 후향적 연구이다. 의무기록을 통한 자료 수집이며, 개인 식별 정보는 익명화 처리되어 수집되므로 대상자의 권리나 복지에 미치는 영향은 없다.")
    h(doc, "6.1 윤리적 연구 수행", 3)
    para(doc, "본 연구는 「생명윤리 및 안전에 관한 법률」에 따라 임상연구의 기본원칙을 준수하며, 헬싱키 선언 및 국제 의학기구협의회(CIOMS)의 국제 윤리 가이드라인에 따라 수행한다.")
    h(doc, "6.2 기관 생명윤리위원회(IRB)", 3)
    para(doc, "연구 개시 전 연구계획서 및 변경사항은 실시기관 IRB의 검토·승인을 받은 후 개시한다. 연구책임자는 IRB에 계획서, 개정본, CRF를 제출하고 승인을 받으며, 제출일·승인일의 최신 목록을 유지한다.")

    # 7. 설계
    h(doc, "7. 연구 설계 및 계획", 1)
    h(doc, "7.1 연구 정의", 3)
    para(doc, "주요 종결점의 정의(사망, 심근경색, 재관류, 뇌졸중, 출혈[BARC], 스텐트 혈전증 등)는 Academic Research Consortium(ARC) 및 BARC 기준을 따른다.")
    h(doc, "7.2 연구 방법", 3)
    para(doc, study['design'])
    para(doc, study['objective'])

    # 8. 수행
    h(doc, "8. 수행", 1)
    para(doc, "대상자에 대한 데이터는 표준 임상 진료에서 이용 가능한 의무기록 정보만 수집한다.")
    h(doc, "8.1 수집 자료", 3)
    for d in study.get('data_collected', []): para(doc, d, bullet=True)

    # 9. 데이터 관리
    h(doc, "9. 데이터 관리", 1)
    h(doc, "9.1 품질 관리", 3)
    para(doc, "데이터 수집에 증례기록서(CRF)를 사용한다.")
    h(doc, "9.2 자료의 비밀성 보장", 3)
    para(doc, "증례기록지에는 개인식별정보를 코드화하여 입력하고, 대상자는 고유 대상자번호로만 식별한다. 수집 자료는 연구자가 직접 익명화하여 개인식별정보를 삭제한 뒤 입력하며, 병록번호·이름 등 식별정보는 데이터베이스 및 결과 발표·출판에서 제외하여 비밀성을 보장한다.")

    # 10. 분석
    h(doc, "10. 자료의 분석", 1)
    h(doc, "10.1 평가항목", 3)
    para(doc, "1차 평가변수: " + study['primary_endpoint'], bold=False)
    if study.get('secondary_endpoints'):
        para(doc, "2차 평가변수:", bold=True)
        for s in study['secondary_endpoints']: para(doc, s, bullet=True)
    h(doc, "10.2 통계적 분석 방법", 3)
    for s in study.get('statistics', []): para(doc, s, bullet=True)

    # 11. 보고서
    h(doc, "11. 보고서", 1)
    para(doc, "연구 종료 후 1년 내에 임상연구 논문 작성을 완료하고 참여 연구기관에 보고한다.")

    # 12. 보관
    h(doc, "12. 연구문서의 보관", 1)
    para(doc, "연구책임자는 임상연구계획서, 개정본, 참여 환자 목록, CRF, 진행보고서를 연구기관 파일에 보관한다. 「생명윤리 및 안전에 관한 법률」에 따라 연구 관련 문서는 연구 종료일로부터 3년간 본원에서 안전하게 보관한다.")

    # 13. 참고문헌
    h(doc, "13. 참고문헌", 1)
    refs, dropped = grounded_refs(study)
    for i, ref in enumerate(refs, 1):
        tag = []
        if ref.get('pmid'): tag.append(f"PMID: {ref['pmid']}")
        if ref.get('doi'): tag.append(f"doi: {ref['doi']}")
        if ref.get('registry'): tag.append(ref['registry'])
        para(doc, f"{i}. {ref['citation']} " + (("(" + "; ".join(tag) + ")") if tag else ""))
    if dropped:
        para(doc, f"[주의] PMID/DOI 미동반으로 제외된 인용 {len(dropped)}건 (Citation Grounding). 검증 후 추가 필요.", bold=True)

    out = os.path.join(outdir, "research_protocol.docx")
    doc.save(out); return out, dropped

# ------------------------------------------------- 요약 [서식 1]
def build_summary(study, outdir):
    doc = Document()
    p = doc.add_paragraph(); r = p.add_run("임상연구 계획서 요약 [서식 1]"); set_font(r, 14, True, BURGUNDY)
    inc = "\n".join(f"{i+1}) {c}" for i, c in enumerate(study['inclusion_criteria']))
    exc = "\n".join(f"{i+1}) {c}" for i, c in enumerate(study['exclusion_criteria']))
    sel = f"[선정기준]\n{inc}\n\n[제외기준]\n{exc}"
    ep = study['primary_endpoint']
    if study.get('secondary_endpoints'):
        ep += "\n\n이차 평가 변수:\n" + "\n".join(study['secondary_endpoints'])
    rows = [
        ("과 제 명", study['title']),
        ("임상연구 목적", study['objective']),
        ("임상연구 참여 기관", study.get('institution','영남대학교병원')),
        ("임상연구 예정 기간", study.get('period','IRB 승인일로부터 2년')),
        ("대상질환", study.get('disease','')),
        ("목표 대상자 수", study.get('target_n','')),
        ("임상시험 의약품/의료기기", study.get('intervention_device','없음')),
        ("선정기준 및 제외기준", sel),
        ("연구디자인", study['design']),
        ("평가항목", ep),
        ("통계분석", "\n".join(study.get('statistics', []))),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2); tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        c0, c1 = tbl.rows[i].cells
        c0.width = Cm(4); c1.width = Cm(12)
        r0 = c0.paragraphs[0].add_run(k); set_font(r0, 10, True)
        r1 = c1.paragraphs[0].add_run(v); set_font(r1, 10)
    out = os.path.join(outdir, "protocol_summary.docx")
    doc.save(out); return out

# --------------------------------- 동의면제 점검표 [서식 41]
def build_waiver(study, outdir):
    doc = Document()
    cw = study.get('consent_waiver', {})
    pi = study['investigators']['pi']
    p = doc.add_paragraph(); r = p.add_run("연구대상자 동의 면제(Waiver of Informed Consent) 점검표 [서식 41]")
    set_font(r, 13, True, BURGUNDY)
    doc.add_paragraph()

    info = doc.add_table(rows=4, cols=1); info.style = 'Table Grid'
    lines = [
        f"■ IRB File No. : {cw.get('irb_file_no','')}",
        f"■ 연구책임자: {pi['name']}",
        f"■ 점 검 자 : {', '.join(cw.get('checkers', []))}",
        "■ 점 검 일 :",
    ]
    for i, ln in enumerate(lines):
        rr = info.rows[i].cells[0].paragraphs[0].add_run(ln); set_font(rr, 10)
    doc.add_paragraph()

    para(doc, "※ 다음은 동의면제가 가능한 경우를 점검합니다.", bold=True)
    criteria = [
        "연구가 대상자에게 미치는 위험이 최소 수준 이하인 경우",
        "동의 면제가 대상자의 권리와 복지에 부정적 영향을 미치지 않는 경우",
        "동의 면제가 없으면 연구 수행이 실질적으로 불가능한 경우",
        "연구 목적상 필요한 경우 동의 면제 후에도 대상자에게 적절한 정보를 추가로 제공할 수 있는 경우",
        "기존 의무기록·자료를 이용하는 후향적 연구로서 개인식별정보가 익명화되어 처리되는 경우",
    ]
    ct = doc.add_table(rows=len(criteria)+1, cols=2); ct.style = 'Table Grid'
    hr = ct.rows[0].cells
    set_font(hr[0].paragraphs[0].add_run("면제사유 점검항목"), 10, True)
    set_font(hr[1].paragraphs[0].add_run("해당(■)"), 10, True)
    for i, c in enumerate(criteria, 1):
        set_font(ct.rows[i].cells[0].paragraphs[0].add_run(c), 10)
        set_font(ct.rows[i].cells[1].paragraphs[0].add_run("■" if cw.get('eligible', True) else "□"), 10)
    doc.add_paragraph()

    verdict = "■ YES        □ NO" if cw.get('eligible', True) else "□ YES        ■ NO"
    para(doc, f"상기 점검항목을 검토한 결과 본 과제는 동의 취득을 면제할 수 있다.  ⇒  {verdict}", bold=True)
    out = os.path.join(outdir, "consent_waiver_checklist.docx")
    doc.save(out); return out

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("study_info")
    ap.add_argument("--outdir", default=".")
    ap.add_argument("--only", choices=["protocol","summary","waiver"], help="단일 문서만 생성")
    a = ap.parse_args()
    os.makedirs(a.outdir, exist_ok=True)
    study = load(a.study_info)
    made = []
    if a.only in (None, "protocol"):
        out, dropped = build_protocol(study, a.outdir); made.append(out)
        if dropped: print(f"[Citation Grounding] PMID/DOI 미동반 인용 {len(dropped)}건 제외됨", file=sys.stderr)
    if a.only in (None, "summary"): made.append(build_summary(study, a.outdir))
    if a.only in (None, "waiver"):  made.append(build_waiver(study, a.outdir))
    for m in made: print("생성:", m)

if __name__ == "__main__":
    main()
